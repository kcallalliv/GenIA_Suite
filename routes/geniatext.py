# routes/alerts.py
from flask import Blueprint, render_template, request, session, redirect, url_for, send_file, jsonify
from google.cloud import bigquery
from google import genai
from google.genai.types import HttpOptions
from functools import wraps
from datetime import datetime
from sqlalchemy.sql import text
from sqlalchemy.orm import aliased
import os
import json
import re
import difflib
import requests
from docx import Document
from bs4 import BeautifulSoup
from docx import Document
from io import BytesIO
from google.api_core.client_options import ClientOptions
from google.cloud import discoveryengine_v1 as discoveryengine
from google import genai
import datetime
import pandas as pd
import traceback # Asume que pandas está importado

client = bigquery.Client(project="prd-claro-mktg-data-storage")
genai_client = genai.Client(
	vertexai=True,
	project="prd-claro-mktg-data-storage",
	location="us-central1",
	http_options=HttpOptions(api_version="v1")
)

geniatext_bp = Blueprint('geniatext_bp', __name__)
@geniatext_bp.route('/genia-text')
def geniatext():
	return render_template('generate.html')
@geniatext_bp.route('/generate-content/select-semrush-keyword', methods=['GET', 'POST'])
def get_semrush_keywords():
	query = """
	SELECT * FROM `prd-claro-mktg-data-storage.project_semantic_seo.semrush_mayo_master`
	"""
	try:
		query_job = client.query(query)
		results = query_job.result()
		keywords_data = list(results) 
		return render_template('generate-select-keyword.html', keywords=keywords_data)
	except Exception as e:
		print("ERROR EJECUTANDO CONSULTA BIGQUERY:")
		traceback.print_exc()  # Muestra el error completo en el log
		return f"Error al ejecutar la consulta de BigQuery: {e}", 500

@geniatext_bp.route('/generate-content/select-search-console-keyword', methods=['GET', 'POST'])
def get_search_console_keywords():
	query = """
		SELECT
			query,
			SUM(impressions) AS total_impressions,
			SUM(clicks) AS total_clicks,
			CAST(AVG(CAST(sum_position AS BIGNUMERIC)) AS INT64) AS avg_position_int
		FROM
			prd-claro-mktg-data-storage.searchconsole.searchdata_url_impression
		WHERE
			-- Excluir keywords de marca y sus variantes/errores
			query NOT LIKE '%cl_ro%'
			AND query NOT LIKE '%calro%'
			AND query NOT LIKE '%claer%'
			AND query NOT LIKE '%clra%'
			AND query NOT LIKE '%mi claro%'
			AND query NOT LIKE '%app claro%'
			AND query NOT LIKE '%claro peru%'
			AND query NOT LIKE '%claro movil%'
			AND query NOT LIKE '%claro hogar%'
			AND query NOT LIKE '%claro empresas%'
			AND query NOT LIKE '%internet claro%'
			AND query NOT LIKE '%telefonia claro%'
			AND query NOT LIKE '%movil claro%'
			AND query NOT LIKE '%fibra optica claro%'
			AND query NOT LIKE '%celular claro%'
			AND query NOT LIKE '%datos claro%'
			AND query NOT LIKE '%plan claro%'
			AND query NOT LIKE '%portabilidad claro%'
			AND query NOT LIKE '%recarga claro%'
			AND query NOT LIKE '%chip claro%'
			AND query NOT LIKE '%tv claro%'
			AND query NOT LIKE '%cable claro%'
			AND query NOT LIKE '%pospago%'
			AND query NOT LIKE '%planes%'
			AND query NOT LIKE '%claeo%'
			AND data_date >= current_date() -30
		GROUP BY
			query
		HAVING
			CAST(AVG(CAST(sum_position AS BIGNUMERIC)) AS INT64) <= 10 and   CAST(AVG(CAST(sum_position AS BIGNUMERIC)) AS INT64) >= 10
		ORDER BY
			total_impressions DESC,
			avg_position_int ASC
		LIMIT 10
	"""
	try:
		query_job = client.query(query)
		results = query_job.result()
		keywords_data = []
		#SELECT Keyword, Search_Volume, Posicion, Intencion 
		#print(f"Error al ejecutar la consulta de BigQuery: {results}")
		#'Keyword', 'Search_Volume', 'Posicion','Intencion'
		for row in results:
			keywords_data.append({
				"Keyword": row.query
			})
		return render_template('generate-select-keyword.html', keywords=keywords_data)
	except Exception as e:
		#return jsonify({"error": f"Error al obtener los datos de BigQuery. Detalle: {e}"}), 500
		print(f"Error al ejecutar la consulta de BigQuery: {e}")
#====Listado Semrush====
@geniatext_bp.route('/generate-content/list-semrush', methods=['GET', 'POST'])
def list_keywords_semrush():
	query = """
	SELECT * FROM `prd-claro-mktg-data-storage.project_semantic_seo.semrush_mayo_master`
	"""
	try:
		query_job = client.query(query)
		results = query_job.result() # Espera a que la consulta termine y obtiene los resultados
		keywords_data = []
		for row in results:
			keywords_data.append({
				"query": row.Keyword,
				"search_volumen": row.Search_Volume,
				"position": row.Posicion,
				"intencion": row.Intencion
			})
		return render_template('generate-listado-semrush-ajax.html', keywords=keywords_data)
	except Exception as e:
		return jsonify({"error": f"Error al obtener los datos de BigQuery. Detalle: {e}"}), 500
@geniatext_bp.route('/generate-content/list-search-console', methods=['GET', 'POST'])
def list_keywords_search_console():
	query = """
	SELECT
		query,
		SUM(impressions) AS total_impressions,
		SUM(clicks) AS total_clicks,
		CAST(AVG(CAST(sum_position AS BIGNUMERIC)) AS INT64) AS avg_position_int
	FROM
		`prd-claro-mktg-data-storage.searchconsole.searchdata_url_impression`
	WHERE
		-- Excluir keywords de marca y sus variantes/errores
		query NOT LIKE '%cl_ro%'
		AND query NOT LIKE '%calro%'
		AND query NOT LIKE '%claer%'
		AND query NOT LIKE '%clra%'
		AND query NOT LIKE '%mi claro%'
		AND query NOT LIKE '%app claro%'
		AND query NOT LIKE '%claro peru%'
		AND query NOT LIKE '%claro movil%'
		AND query NOT LIKE '%claro hogar%'
		AND query NOT LIKE '%claro empresas%'
		AND query NOT LIKE '%internet claro%'
		AND query NOT LIKE '%telefonia claro%'
		AND query NOT LIKE '%movil claro%'
		AND query NOT LIKE '%fibra optica claro%'
		AND query NOT LIKE '%celular claro%'
		AND query NOT LIKE '%datos claro%'
		AND query NOT LIKE '%plan claro%'
		AND query NOT LIKE '%portabilidad claro%'
		AND query NOT LIKE '%recarga claro%'
		AND query NOT LIKE '%chip claro%'
		AND query NOT LIKE '%tv claro%'
		AND query NOT LIKE '%cable claro%'
		AND query NOT LIKE '%pospago%'
		AND query NOT LIKE '%planes%'
		AND query NOT LIKE '%claeo%'
		AND data_date >= current_date() -30
	GROUP BY
		query
	HAVING
		CAST(AVG(CAST(sum_position AS BIGNUMERIC)) AS INT64) <= 10
	ORDER BY
		total_impressions DESC,
		avg_position_int ASC
	LIMIT 10
	"""
	try:
		# Ejecutar la consulta en BigQuery
		query_job = client.query(query)
		results = query_job.result() # Espera a que la consulta termine y obtiene los resultados
		keywords_data = []
		for row in results:
			keywords_data.append({
				"query": row.query,
				"total_impressions": row.total_impressions,
				"total_clicks": row.total_clicks,
				"avg_position_int": row.avg_position_int
			})
		return render_template('generate-listado-ajax.html', keywords=keywords_data)
	except Exception as e:
		# Manejo de cualquier otro error durante la ejecución de la consulta
		#print(f"Error al ejecutar la consulta de BigQuery: {e}")
		return jsonify({"error": f"Error al obtener los datos de BigQuery. Detalle: {e}"}), 500
		#return render_template('generate-listado-ajax.html', keywords=keywords_data)
#json
@geniatext_bp.route('/generate-content/select-tipo-contenido', methods=['GET', 'POST'])
def get_tipo_contenido():
	base_dir = os.path.dirname(os.path.abspath(__file__))
	json_file_path = os.path.join(base_dir, '..', 'static', 'json', 'json_tipo-contenido.json')
	# ¡INICIALIZA la lista aquí!
	contenido_list = [] 
	try:
		with open(json_file_path, 'r', encoding='utf-8') as f:
			data = json.load(f)
		# Listar y construir la lista de diccionarios
		for tipo, detalles in data.items():
			contenido_item = {
				"nombre_tipo": tipo,
				"intencion": detalles.get("intencion", "N/A"),
				"descripcion": detalles.get("descripcion", "N/A"),
				"ejemplo": detalles.get("ejemplo", "N/A")
			}
			contenido_list.append(contenido_item)
		return render_template('generate-select-tipo-contenido.html', tipo_contenido=contenido_list)
	except FileNotFoundError:
		print(f"Error: El archivo '{json_file_path}' no se encontró.")
		return "Error: Archivo de tipos de contenido no encontrado.", 500 
	except json.JSONDecodeError:
		print(f"Error: No se pudo decodificar el archivo JSON '{json_file_path}'. Asegúrate de que sea un JSON válido.")
		return "Error: Formato de archivo de tipos de contenido inválido.", 500
	except Exception as e:
		print(f"Ocurrió un error inesperado al cargar el JSON: {e}")
		return f"Error interno del servidor: {e}", 500
@geniatext_bp.route('/generate-content/slide-filter', methods=['GET', 'POST'])
def get_slider_filter():
	return render_template('generate-slide-filter.html')

#---IA Generate---
def get_json_tipo_contenido():
	base_dir = os.path.dirname(os.path.abspath(__file__))
	json_file_path = os.path.join(base_dir, '..', 'static', 'json', 'json_tipo-contenido.json')
	contenido_list = []
	try:
		with open(json_file_path, 'r', encoding='utf-8') as f:
			data = json.load(f) # Carga el JSON tal cual
		return data # Retorna el diccionario directamente
	except FileNotFoundError:
		print(f"Error: El archivo JSON no se encontró en la ruta: {json_file_path}")
		return []
	except json.JSONDecodeError:
		print(f"Error: No se pudo decodificar el archivo JSON '{json_file_path}'. Asegúrate de que sea un JSON válido.")
		return []
	except Exception as e:
		print(f"Ocurrió un error inesperado al cargar el JSON: {e}")
		return []
def get_urls_existentes():
	query = """
	SELECT * FROM prd-claro-mktg-data-storage.project_semantic_seo.hablando_claro_url
	"""
	return client.query(query).to_dataframe()
import re
import pandas as pd 
# Asume que genai_client ya está configurado y accesible

def elegir_urls_relevantes_con_gemini(keyword, urls_df):
	if urls_df.empty:
		return "No se encontraron URLs similares para sugerir."
	if 'page_url' not in urls_df.columns:
		return "Error interno: columna 'page_url' no encontrada."

	urls_texto = "\n".join([f"- {row['page_url']}" for _, row in urls_df.iterrows()])
	prompt = f"""
	Eres un experto en SEO. Tu tarea es identificar las 3 URLs más relevantes para una correcta implementación de enlaces internos.

	**Keyword Principal del Artículo:** '{keyword}'

	**Lista de URLs DISPONIBLES (SOLO puedes usar URLs de esta lista, NO inventes ni modifiques):**
	{urls_texto}

	**INSTRUCCIONES DE SALIDA ESTRICTAS:**
	Genera SOLO un listado numérico de 3 URLs. Cada URL debe estar en una nueva línea, precedida por su número y un punto, y sin ningún texto adicional o explicaciones.

	EJEMPLO DE FORMATO DE SALIDA (EXACTO):
	1. https://www.ejemplo.com/primera-url
	2. https://www.ejemplo.com/segunda-url
	3. https://www.ejemplo.com/tercera-url

	Asegúrate de que las URLs seleccionadas sean las más lógicamente relacionadas con la keyword principal.
	"""
	try:
		response = genai_client.models.generate_content(
			model="gemini-2.0-flash-001",
			contents=prompt
		)
		gemini_raw_response = response.text.strip()
		url_pattern = re.compile(r'^\d+\.\s*(https?://\S+)$', re.MULTILINE)
		suggested_urls = url_pattern.findall(gemini_raw_response)
		return suggested_urls if suggested_urls else "No se encontraron URLs relevantes para sugerir."
	except Exception as e:
		print(f"ERROR: elegir_urls_relevantes_con_gemini falló con una excepción: {e}")
		return "Error al generar sugerencias de URLs."
def get_urls_similares_por_keyword(keyword, urls_df):
	print(f"INFO: Buscando URLs similares para keyword: '{keyword}'")
	textos = urls_df["page_url"].tolist()
	# Asegúrate de que 'textos' no esté vacío antes de llamar a difflib.get_close_matches
	if not textos:
		print("ADVERTENCIA: No hay URLs en el DataFrame para buscar coincidencias.")
		return pd.DataFrame()
	matches = difflib.get_close_matches(keyword, textos, n=10, cutoff=0.2)
	print(f"INFO: Coincidencias de URLs encontradas: {matches}")
	return urls_df[urls_df["page_url"].isin(matches)].reset_index(drop=True)
def obtener_contexto_vertex(keyword):
	print(f"INFO: Obtener contexto de Vertex AI Search para keyword: '{keyword}'")
	project_id = "prd-claro-mktg-data-storage"
	location = "global"
	engine_id = "claro-peru_1748977843565"

	try:
		client_options = ClientOptions(api_endpoint=f"{location}-discoveryengine.googleapis.com") if location != "global" else None
		client = discoveryengine.ConversationalSearchServiceClient(client_options=client_options)

		serving_config = f"projects/{project_id}/locations/{location}/collections/default_collection/engines/{engine_id}/servingConfigs/default_serving_config"

		answer_generation_spec = discoveryengine.AnswerQueryRequest.AnswerGenerationSpec(
			ignore_adversarial_query=True,
			ignore_non_answer_seeking_query=True,
			ignore_low_relevant_content=False,
			model_spec=discoveryengine.AnswerQueryRequest.AnswerGenerationSpec.ModelSpec(model_version="stable"),
			prompt_spec=discoveryengine.AnswerQueryRequest.AnswerGenerationSpec.PromptSpec(
				preamble="Usa la información más relevante de tus fuentes indexadas, brinda el contexto necesario para usar info dentro de un articulo de Claro, comparte 3 urls relacionadas para que puedan buscar informacion mas detallada los usuarios."
			),
			include_citations=True,
			answer_language_code="es",
		)

		request = discoveryengine.AnswerQueryRequest(
			serving_config=serving_config,
			query=discoveryengine.Query(text=keyword),
			session=None,
			answer_generation_spec=answer_generation_spec,
		)

		print(f"INFO: Enviando solicitud a Vertex AI Search para keyword: '{keyword}'")
		response = client.answer_query(request)
		print(f"INFO: Respuesta de Vertex AI Search recibida. ¿Contiene respuesta?: {bool(response.answer)}")
		return response.answer.answer_text if response.answer else ""
	except Exception as e:
		st.error(f"ERROR al obtener contexto de Vertex AI Search: {e}")
		print(f"ERROR: obtener_contexto_vertex falló con: {e}")
		return f"Error al obtener contexto: {e}"
def generar_articulo(keyword, tipo_contenido_seleccionado, tipo_contenido_libre, tema_especifico):
	print(f"INFO: Iniciando generación de artículo para keyword: '{keyword}', tipo: '{tipo_contenido_seleccionado}' (libre: '{tipo_contenido_libre}') y tema: '{tema_especifico}'")
	# Llamar Tipo contenido
	TIPOS_CONTENIDO = get_json_tipo_contenido()
	# Paso 1: obtener contexto desde documentos indexados
	contexto_vertex = obtener_contexto_vertex(keyword)
	print(f"INFO: Contexto de Vertex obtenido. Longitud: {len(contexto_vertex)} caracteres.")

	# Paso 2: obtener URLs sugeridas y tipo de contenido
	urls_df = get_urls_existentes()
	similares_df = get_urls_similares_por_keyword(keyword, urls_df)
	sugerencias_urls = elegir_urls_relevantes_con_gemini(keyword, similares_df)
	print(f"INFO: URLs relacionadas sugeridas: {sugerencias_urls}")

	# Determinar la intención y descripción final
	if tipo_contenido_seleccionado == "Otro (especificar)":
		intencion = f"Crear un contenido de tipo: '{tipo_contenido_libre}'"
		descripcion_tipo = f"El contenido se enfocará en el tipo '{tipo_contenido_libre}', adaptándose a la keyword proporcionada."
		tipo_final_para_prompt = tipo_contenido_libre
	else:
		detalles_tipo = TIPOS_CONTENIDO.get(tipo_contenido_seleccionado, {})
		intencion = detalles_tipo.get("intencion", "Saber sobre el tema")
		descripcion_tipo = detalles_tipo.get("descripcion", "")
		tipo_final_para_prompt = tipo_contenido_seleccionado

	print(f"INFO: Detalles del tipo de contenido final: Intención='{intencion}', Descripción='{descripcion_tipo}'")

	# Paso 3: construir el prompt incluyendo el contexto Vertex y el tema específico
	prompt = (
		"Eres un experto en SEO y Content Marketing con más de 20 años de experiencia, especializado en crear artículos de alta calidad para portales web del sector telecomunicaciones y negocios en Perú. "
		"Tu contenido está enfocado en mejorar el posicionamiento SEO, aumentar el tiempo de permanencia, el scroll y la navegación interna. "
		"No redirijas a tiendas comerciales. Todo debe ser útil y relevante para el usuario peruano.\n\n"

		f"**Keyword Principal:** {keyword}\n" # Explicitamos la keyword aquí
		f"**Tipo de Contenido Solicitado:** {tipo_final_para_prompt}\n"
		f"**Intención de Búsqueda del Usuario:** {intencion}\n"
		f"**Descripción del Tipo de Contenido:** {descripcion_tipo}\n"
	)

	# Añadir el tema específico al prompt si se proporcionó
	if tema_especifico:
		prompt += f"**Tema Específico del Artículo:** {tema_especifico}\n\n" # <--- Nuevo en el prompt
		prompt += "Asegúrate de que el contenido se centre exclusivamente en este tema específico.\n\n"
	else:
		prompt += "\n" # Para mantener el formato si no hay tema específico

	prompt += (
		f"**Contexto Relevante desde Documentos Indexados con Urls internas disponibles (Vertex AI Search):**\n{contexto_vertex}\n\n"

		"Antes de redactar:\n"
		"- Analiza a profundidad la keyword y, si aplica, el tema específico.\n" # <--- Pequeño ajuste aquí
		"- Investiga temas, subtemas y preguntas relacionadas más buscadas en Google (autocomplete).\n"
		"- Integra esas búsquedas dentro del contenido de forma natural.\n"
		"- Usa redacción fluida, clara y natural, sin repetir innecesariamente la keyword.\n"
		"- Incluye sinónimos o términos relacionados estratégicamente.\n"
		"- Relaciona naturalmente conceptos con otros artículos o secciones del portal cuando sea posible.\n"
		"- Sugiere al lector, de manera útil y sutil, visitar páginas específicas del sitio web si quiere ampliar el tema.\n"
		"- Solo puedes redirigir a las URLs proporcionadas arriba, si el contenido lo permite de forma natural.\n"
		"- Cuando enlaces a otra URL, hazlo con frases como: 'como explicamos en nuestro artículo sobre...', 'puedes ver más en...', 'también te puede interesar...'.\n"
		"- Solo puedes mencionar la marca 'Claro' un máximo de dos veces en todo el artículo.\n\n"

		f"Ahora genera el contenido para la keyword: '**{keyword}**'."
	)
	
	# Ajuste para que el prompt final se adapte si hay tema específico
	if tema_especifico:
		prompt += f" El artículo debe tratar sobre el tema: '{tema_especifico}'. El texto debe ser claro, útil, informativo y redactado para un usuario peruano.\n\n"
	else:
		prompt += " El texto debe ser claro, útil, informativo y redactado para un usuario peruano.\n\n"
		
	prompt += (
		"Formato de salida obligatorio. Utiliza las etiquetas de marcador exactamente como se especifican:\n"
		"[[TITULO]]\n"
		"**Título principal del artículo**\n\n"
		"[[INTRODUCCION]]\n"
		"**Introducción:** Un párrafo de máximo 4 líneas, donde la keyword principal debe estar en negrita (ej. **tu keyword**).\n\n"
		"[[SUBTITULO]]\n"
		"**Primer subtítulo del artículo (siempre en negrita, ej. ****Tu Subtítulo****)**\n\n" # <-- ¡Cambio aquí!
		"[[CUERPO]]\n"
		"**Cuerpo del artículo:** Desarrolla el tema entre 600 y 1200 palabras. Todos los subtítulos dentro del cuerpo del artículo también deben estar en negrita (usando ****). Cada sección debe estar separada por saltos de línea (párrafos vacíos entre ellas).\n\n" # <-- ¡Cambio aquí!
		"[[META_TITLE]]\n"
		"**Meta title:** Máximo 65 caracteres, debe terminar con '| Hablando Claro'.\n\n"
		"[[META_DESCRIPTION]]\n"
		"**Meta description:** Máximo 160 caracteres.\n\n"
		"Asegúrate de que cada sección esté claramente separada por las etiquetas y saltos de línea (párrafos vacíos entre ellas).\n"
		"No brindes más explicaciones, preguntas frecuentes, ni menciones a la competencia. Céntrate únicamente en generar el contenido solicitado en el formato especificado."
	)

	try:
		print(f"INFO: Enviando prompt final a Gemini para generación de artículo. Prompt (parcial): {prompt[:500]}...")
		response = genai_client.models.generate_content(
			model="gemini-2.0-flash-001",
			contents=prompt
		)
		print(f"INFO: Artículo generado por Gemini. Longitud: {len(response.text)} caracteres.")
		return response.text, sugerencias_urls
	except Exception as e:
		st.error(f"ERROR al generar el artículo con Gemini: {e}")
		print(f"ERROR: generar_articulo (llamada a Gemini) falló con: {e}")
		return f"Error al generar el artículo: {e}", sugerencias_urls
import re
def extract_sections_from_gemini_text(gemini_text):
	parsed_sections = {}
	markers = [
		"[[TITULO]]", "[[INTRODUCCION]]", "[[SUBTITULO]]",
		"[[CUERPO]]", "[[META_TITLE]]", "[[META_DESCRIPTION]]"
	]
	split_pattern = r'(' + '|'.join(re.escape(m) for m in markers) + r')'
	parts = re.split(split_pattern, gemini_text)

	current_marker_key = None
	buffer = []

	for part in parts:
		part = part.strip()
		if not part: continue

		if part in markers:
			if current_marker_key and buffer:
				parsed_sections[current_marker_key] = "\n".join(buffer).strip()
			current_marker_key = part.replace("[[", "").replace("]]", "").lower()
			buffer = []
		else:
			buffer.append(part)
	
	if current_marker_key and buffer:
		parsed_sections[current_marker_key] = "\n".join(buffer).strip()
	
	return parsed_sections
def parse_markdown_to_html(content_to_parse_with_markers):
	html_content_parts = []
	sections_for_html = extract_sections_from_gemini_text(content_to_parse_with_markers)

	# Patrones para enlaces Markdown y URLs planas
	markdown_link_pattern = re.compile(r'\[(.*?)\]\((https?://[^\s)]+)\)')
	plain_url_pattern = re.compile(r'(https?://[^\s<>"]+)(?<![.,;:])') # Adjusted to capture URL cleanly

	# Helper para convertir texto con enlaces y negritas, pero sin añadir <p> o <br> de inmediato
	def apply_inline_formatting(text):
		# Primero, procesa enlaces Markdown [texto](url)
		text = markdown_link_pattern.sub(r'<a href="\2" target="_blank">\1</a>', text)
		
		# Luego, procesa URLs planas que no estén ya dentro de un <a> tag
		# Lookbehind para asegurar que no precede a ">" (fin de tag) o ya es parte de un href="..."
		text = re.sub(r'(?<!href=["\'])(?<!>)(https?://[^\s<>"]+)(?<![.,;:])', r'<a href="\1" target="_blank">\1</a>', text)
		
		# Procesa negritas **texto**
		text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)
		return text

	# --- Título ---
	title_content = sections_for_html.get('titulo', '').strip()
	if title_content:
		# El título original podría tener ** alrededor, los quitamos antes de aplicar formato.
		# Asegúrate de que el título no tenga ** iniciales y finales si es un solo título
		title_content_clean = re.sub(r'^\*\*(.*?)\*\*$', r'\1', title_content).strip()
		html_content_parts.append(f'<div class="title">{apply_inline_formatting(title_content_clean)}</div>')

	# --- Introducción ---
	intro_content = sections_for_html.get('introduccion', '').strip()
	if intro_content:
		# Aplicar formato de enlaces/negritas, luego manejar saltos de línea y p-tag
		processed_intro = apply_inline_formatting(intro_content)
		# Reemplazar saltos de línea por <br> solo si no hay ya tags HTML que manejen el espaciado
		processed_intro = processed_intro.replace('\n', '<br>')
		html_content_parts.append(f'<p class="introduction">{processed_intro}</p>')

	# --- Subtítulo principal ---
	subtitle_content = sections_for_html.get('subtitulo', '').strip()
	if subtitle_content:
		# Quitar negritas si el subtítulo completo está envuelto en ellas
		subtitle_content_clean = re.sub(r'^\*\*(.*?)\*\*$', r'\1', subtitle_content).strip()
		# Aplicar formato de enlaces/negritas
		html_content_parts.append(f'<p class="subtitle"><strong>{apply_inline_formatting(subtitle_content_clean)}</strong></p>')

	# --- Cuerpo ---
	body_content = sections_for_html.get('cuerpo', '').strip()
	if body_content:
		processed_body_blocks = []
		# Dividir el cuerpo en bloques por saltos de línea dobles
		# Esto ayuda a tratar párrafos y subtítulos internos de forma separada
		body_blocks = re.split(r'\n{2,}', body_content.strip())
		
		for block in body_blocks:
			block = block.strip()
			if not block:
				continue

			# Check for internal subtitles like **Este Es Un Subtítulo Interno**
			if block.startswith('**') and block.endswith('**') and len(block) > 4:
				# Quita los asteriscos y aplica formato, luego envuelve en su propio párrafo de subtítulo
				subtitle_text = block.replace('**', '').strip()
				processed_body_blocks.append(f'<p class="body-subtitle"><strong>{apply_inline_formatting(subtitle_text)}</strong></p>')
			else:
				# Si no es un subtítulo, es un párrafo normal o una URL dedicada
				# Aplicar formato inline (enlaces, negritas)
				formatted_block = apply_inline_formatting(block)
				
				# Check if the entire block is just a plain URL or markdown link, make it a <p class="link">
				if re.fullmatch(r'^\s*(<a href="https?://[^\s<>"]+" target="_blank">.*?</a>)\s*$', formatted_block, re.DOTALL):
					processed_body_blocks.append(f'<p class="link">{formatted_block}</p>')
				else:
					# Si no es un enlace dedicado, reemplazar saltos de línea por <br>
					# y envolver en un <p> normal
					formatted_block_with_br = formatted_block.replace('\n', '<br>')
					processed_body_blocks.append(f'<p>{formatted_block_with_br}</p>')

		html_content_parts.append("".join(processed_body_blocks))

	# --- URLs Relacionadas (si se extraen como una sección separada) ---
	urls_related_content = sections_for_html.get('urls_relacionadas', '').strip()
	if urls_related_content:
		document.add_heading("URLs Relacionadas:", level=3)
		url_lines = urls_related_content.split('\n')
		for url_item in url_lines:
			url_item = url_item.strip()
			if not url_item: continue

			# Procesa tanto enlaces Markdown como URLs planas dentro de la sección
			processed_url_item = apply_inline_formatting(url_item)
			
			# Envuelve cada URL procesada en un párrafo con la clase 'link'
			html_content_parts.append(f'<p class="link">{processed_url_item}</p>')
	
	# --- Meta Title y Meta Description (si se extraen como secciones separadas) ---
	meta_title_content = sections_for_html.get('meta_title', '').strip()
	if meta_title_content:
		html_content_parts.append(f'')

	meta_description_content = sections_for_html.get('meta_description', '').strip()
	if meta_description_content:
		html_content_parts.append(f'')
	
	return "\n\n".join(html_content_parts)

"""
@geniatext_bp.route('/generate-content/slide-result', methods=['GET', 'POST'])
def get_slider_result():
	keyword = "internet"
	tipo_contenido_seleccionado = "Semrush"
	tipo_contenido_libre = "medidor de velocidad"
	tema_especifico = "internet"
	#generar_articulo(keyword, tipo_contenido_seleccionado, tipo_contenido_libre, tema_especifico)
	return render_template('generate-slide-result.html')

#@geniatext_bp.route('/generate-content/article', methods=['GET', 'POST'])
"""
#Crear Documento Word
def crear_docx(texto_plain, name_file):
	document = Document()

	sections = extract_sections_from_gemini_text(texto_plain)

	inferred_keyword = name_file.replace(".docx", "").replace("_", " ").title()
	document.add_heading(f"Documento Generado: {inferred_keyword}", level=1)
	document.add_paragraph()

	if 'titulo' in sections and sections['titulo']:
		title_text = sections['titulo'].replace('**', '').strip()
		document.add_heading(title_text, level=2)
	
	if 'introduccion' in sections and sections['introduccion']:
		intro_text = sections['introduccion'].replace('**', '').strip()
		document.add_paragraph("Introducción:")
		document.add_paragraph(intro_text)
		document.add_paragraph()

	if 'subtitulo' in sections and sections['subtitulo']:
		subtitle_text = sections['subtitulo'].replace('**', '').strip()
		document.add_heading(subtitle_text, level=3)
		document.add_paragraph()

	if 'cuerpo' in sections and sections['cuerpo']:
		body_text = sections['cuerpo']
		body_blocks = re.split(r'\n{2,}', body_text.strip())
		
		for block in body_blocks:
			block = block.strip()
			if not block:
				continue
			
			if block.startswith('**') and block.endswith('**') and len(block) > 4:
				document.add_heading(block.replace('**', '').strip(), level=4)
			else:
				p = document.add_paragraph()
				parts = re.split(r'(\*\*[^*]+\*\*)', block)
				for part in parts:
					if part.startswith('**') and part.endswith('**'):
						p.add_run(part[2:-2]).bold = True
					else:
						p.add_run(part)
		document.add_paragraph()

	if 'urls_relacionadas' in sections and sections['urls_relacionadas']:
		document.add_heading("URLs Relacionadas:", level=3)
		url_lines = sections['urls_relacionadas'].split('\n')
		for url in url_lines:
			url = url.strip()
			if url and (url.startswith("http://") or url.startswith("https://")):
				p = document.add_paragraph()
				p.add_run("• ").bold = True
				p.add_run(url, style='Hyperlink')
			elif url:
				document.add_paragraph(f"• {url}")
		document.add_paragraph()

	if ('meta_title' in sections and sections['meta_title']) or \
	   ('meta_description' in sections and sections['meta_description']):
		document.add_heading("Datos SEO:", level=3)
		if 'meta_title' in sections and sections['meta_title']:
			document.add_paragraph(f"Meta Title: {sections['meta_title'].strip()}")
		if 'meta_description' in sections and sections['meta_description']:
			document.add_paragraph(f"Meta Description: {sections['meta_description'].strip()}")
		document.add_paragraph()

	buffer = BytesIO()
	document.save(buffer)
	buffer.seek(0)
	return buffer

@geniatext_bp.route('/generate-content/article-result', methods=['GET', 'POST'])
def get_slider_article():
	if request.method in ['GET', 'POST']:
		vfuente = request.values.get('fuente')
		vkeyword = request.values.get('keyword')
		vtipo_articulo = request.values.get('tipo_articulo')
		varticle = request.values.get('article')
		vcustom_keyword = request.values.get('custom_keyword')
		full_gemini_text, sugerencias_urls = generar_articulo(vkeyword, vtipo_articulo, vcustom_keyword, varticle)
		parsed_plain_sections = extract_sections_from_gemini_text(full_gemini_text)
		meta_title = parsed_plain_sections.get('meta_title', 'Meta title no encontrado').strip()
		meta_description = parsed_plain_sections.get('meta_description', 'Meta description no encontrada').strip()
		article_content_for_html_parsing = ""
		for sec_key in ['titulo', 'introduccion', 'subtitulo', 'cuerpo']:
			if sec_key in parsed_plain_sections:
				article_content_for_html_parsing += f"[[{sec_key.upper()}]]\n" + parsed_plain_sections[sec_key] + "\n\n"
		final_html_content = parse_markdown_to_html(article_content_for_html_parsing)
		session['article_content_for_docx'] = article_content_for_html_parsing
		session['article_keyword_for_docx'] = vkeyword # También guarda la keyword para el nombre del archivo
		articulo_data = {
			"content": final_html_content,
			"meta_title": meta_title,
			"meta_description": meta_description
		}
	return render_template('generate-slide-result.html', articulo=articulo_data, urls_sugeridas=sugerencias_urls)
# --- Nueva ruta para la descarga del documento Word ---
@geniatext_bp.route('/download-generated-article', methods=['GET'])
def download_generated_article():
	# Recuperar el contenido guardado en la sesión
	texto_plain_for_docx = session.pop('article_content_for_docx', None)
	keyword_for_docx = session.pop('article_keyword_for_docx', None)
	if not texto_plain_for_docx or not keyword_for_docx:
		# Si no hay contenido en la sesión, redirigir o mostrar un error
		return "No hay artículo para descargar. Por favor, genera un artículo primero.", 400
	name_file = f"{keyword_for_docx}_articulo.docx"
	# Llamar a crear_docx con los datos recuperados de la sesión
	docx_buffer = crear_docx(texto_plain_for_docx, name_file)
	# Retornar el archivo como una respuesta de Flask
	return send_file(
		docx_buffer,
		mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
		as_attachment=True,
		download_name=name_file
	)