# routes/alerts.py
from flask import Blueprint, render_template, request, session, redirect, url_for, send_file, jsonify
from google.cloud import bigquery
from google import genai
from google.genai.types import HttpOptions
from google.api_core.client_options import ClientOptions
from google.cloud import storage
from google.cloud import discoveryengine_v1 as discoveryengine
from google.genai import types
from functools import wraps
from datetime import datetime
from datetime import timedelta
import datetime
import sys
import os
import json
import re
import hashlib
import difflib
import requests
import random
import string
from io import BytesIO
import mimetypes
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import pandas as pd
import traceback # Asume que pandas está importado
import uuid
# Importaciones necesarias para la creación manual del hipervínculo:
import docx.opc.constants
import docx.oxml.shared
import docx.text.run
from docx.oxml.ns import qn
from docx.enum.dml import MSO_THEME_COLOR
from docx.shared import RGBColor, Pt 
#llamar a config
from routes.models import db, Assets, Keywords, Imagesia, Configuracion
from routes.config.geniaconfig import bq_client, bucket_name, storage_client, bq_table_config, genai_client, bq_table_brands, validar_sesion, fechaActual, generarCodigo

client = bigquery.Client(project="prd-claro-mktg-data-storage")

cltarticle_bp = Blueprint('cltarticle_bp', __name__, url_prefix='/brand/<string:brand_id>/proyectos/<string:proyecto_id>/colecciones/article/')

@cltarticle_bp.route('/')
@validar_sesion
def article_main(brand_id,proyecto_id):
	key_id = request.values.get('keyid')
	pid = proyecto_id
	data = Keywords.query.filter_by(keyword_id=key_id).first()
	return render_template('sections/colecciones/article/generate-main.html',data=data,pid=pid)

@cltarticle_bp.route('images', methods=['GET', 'POST'])
@validar_sesion
def article_images(brand_id,proyecto_id):
	pid = request.values.get('pid')
	lista = listarImages(pid)
	return render_template('sections/colecciones/article/images.html',lista=lista,generar_url_firmada=generar_url_firmada)

@cltarticle_bp.route('generated-title', methods=['GET', 'POST'])
def generaTitleIA(brand_id, proyecto_id):
	tipo_contenido = request.values.get('tipo')
	tema = request.values.get('tema')
	keyword = request.values.get('keyword')

	try:
		tipo_final_para_prompt = tipo_contenido
		# --- Construcción del Prompt (Instrucción simplificada) ---
		prompt = (
			"Eres un experto en SEO y Content Marketing con más de 20 años de experiencia, especializado en crear títulos de alto impacto y orientados al usuario peruano. "
			"Tu único objetivo es generar un título que sea atractivo, **que contenga obligatoriamente la Keyword Principal** y utilice el Tipo de Contenido y el Tema Específico como **contexto y guía creativa**.\n\n"
			
			# --- DELIMITADOR DE DATOS DE ENTRADA ---
			"**DATOS DE ENTRADA:**\n"
			f"Keyword Principal: {keyword}\n"
			f"Tipo de Contenido: {tipo_final_para_prompt}\n"
			f"Tema Específico: {tema}\n"
			"**FIN DE DATOS DE ENTRADA.**\n\n"
			
			# --- INSTRUCCIONES ESTRICTAS DE SALIDA ---
			"**INSTRUCCIÓN CLAVE:** Genera **SOLO** el texto del título principal. La salida debe ser una **única línea de texto**.\n"
			"**USO DE KEYWORD:** El título debe incluir la Keyword Principal, pero **NO es necesario que aparezca al inicio**. Intégrala de forma natural en cualquier parte de la frase que tenga sentido semántico.\n"
			"**FORMATO Y ESTRUCTURA:** El título debe ser una **frase fluida y natural**. **ESTÁ TERMINANTEMENTE PROHIBIDO** usar la estructura 'Tema: Título'. **SUSTITUCIÓN OBLIGATORIA:** Si vas a usar dos puntos (:) para separar una idea, **USA UNA COMA (,)** en su lugar. La coma es tu separador principal.\n"
			"**ESTILO DE TÍTULO:** Genera títulos informativos o de lista, **EVITANDO** el formato de pregunta directa (signos ¿?) y frases genéricas. Céntrate en el valor del tema.\n"
			"**ORTOGRAFÍA Y CAPITALIZACIÓN:** El título debe seguir las reglas ortográficas del español. Usa mayúscula **solo** al inicio de la frase y en nombres propios.\n"
			"**PROHIBICIÓN ABSOLUTA:** NO incluyas el carácter de **dos puntos (:)** bajo ninguna circunstancia. Si el título generado lo tiene, elimínalo y reestructura la frase.\n\n"
			
			# --- INSTRUCCIÓN FINAL DE ANIQUILACIÓN DE TEXTO DE RELLENO ---
			"**NO** incluyas ninguna explicación, introducción, comentario o frase de verificación de datos. **SOLO** el título final."
		)
		# --- Llamada a Gemini ---
		# Asegúrate de que 'genai_client' esté definido e inicializado correctamente antes de esta función
		# response = genai_client.models.generate_content(...) 
		
		# Simulando la respuesta para que el código sea funcional en este contexto
		# contenido_generado = "El Mejor título sobre **hogar** y **celulares** que contenga la palabra **claro**"
		
		# Usando la línea original, asumiendo que 'genai_client' está definido
		response = genai_client.models.generate_content(
			model="gemini-2.0-flash-001",
			contents=prompt
		)
		contenido_generado = response.text

		# --- Extracción y Retorno del Título (El parser) ---
		# Si pedimos negritas, buscamos el patrón: ** (Texto del título) **
		match = re.search(r'\*\*([^*]+)\*\*', contenido_generado)
		result_title = "" # Inicializar variable

		if match:
			result_title = match.group(1).strip()
		else:
			# Si la IA no usa negritas, devolvemos el texto completo y limpio
			result_title = contenido_generado.strip()
		
		# Verificación de longitud para manejo de errores de formato
		if len(result_title) < 5:
			# Si es muy corto, aún lo devolvemos, pero se puede considerar un error si es crítico
			pass

		return jsonify({
			"success": True,
			"title": result_title
		})
		
	except Exception as e:
		error_msg = f"ERROR: Fallo crítico en la función generaTitleIA: {e}"
		
		# 🚨 SOLUCIÓN: Usamos el logger estándar en lugar de st.error
		logging.error(error_msg) 
		
		result_title = f"Error de ejecución: {e}"
		return jsonify({
			"success": False,
			"title": result_title
		})
@cltarticle_bp.route('article-content', methods=['GET', 'POST'])
def generaContenidoIA(brand_id, proyecto_id):
	# --- 1. Datos de Entrada ---
	tipo_contenido = request.values.get('tipo')
	tema = request.values.get('tema')
	title = request.values.get('titulo')
	keyword = request.values.get('keyword')

	# --- 2. Configuración y Contexto ---
	json_contenido = get_json_tipo_contenido()
	detalles_tipo = json_contenido.get(tipo_contenido, {})
	intencion = detalles_tipo.get("intencion", "Contratar o mejorar servicio")
	
	contexto_vertex = obtener_contexto_vertex(keyword)
	urls_df = get_urls_existentes()
	similares_df = get_urls_similares_por_keyword(keyword, urls_df)
	sugerencias_urls = elegir_urls_relevantes_con_gemini(keyword, similares_df)

	# --- 3. Lógica de Selección de Producto (Python) ---
	# Esto decide de qué hablar ANTES de llamar a la IA para asegurar el enfoque comercial.
	kw_lower = keyword.lower()
	instruccion_producto = ""

	if any(x in kw_lower for x in ['internet', 'wifi', 'velocidad', 'hogar', 'fibra', 'lento', 'router']):
		instruccion_producto = (
			"**ENFOQUE DE PRODUCTO (OBLIGATORIO):**\n"
			"El tema central debe ser la **Fibra Óptica de Claro**.\n"
			"- Destaca: Estabilidad, velocidad simétrica (subida=bajada) y cobertura.\n"
			"- Argumento: 'La red más estable para trabajar, estudiar y jugar sin lag'."
		)
	elif any(x in kw_lower for x in ['celular', 'movil', 'chip', 'datos', 'prepago', 'recarga', 'megas']):
		instruccion_producto = (
			"**ENFOQUE DE PRODUCTO (OBLIGATORIO):**\n"
			"El tema central debe ser **Claro Prepago** (énfasis en economía/recargas) o **Planes Postpago Max**.\n"
			"- Destaca: Cobertura nacional, apps ilimitadas y beneficios exclusivos."
		)
	elif any(x in kw_lower for x in ['tv', 'cine', 'pelicula', 'series', 'streaming', 'cable']):
		instruccion_producto = (
			"**ENFOQUE DE PRODUCTO (OBLIGATORIO):**\n"
			"El tema central debe ser **Claro TV+** y los beneficios de **Claro Video**.\n"
			"- Destaca: Que los planes incluyen acceso a contenido y Paramount+ (según vigencia)."
		)
	else:
		# Default si la keyword es muy genérica
		instruccion_producto = (
			"**ENFOQUE DE PRODUCTO:**\n"
			"Menciona sutilmente cómo la tecnología y conectividad de **Claro** resuelven la necesidad del usuario."
		)

	try:
		template = """<h1>[titulo]</h1>
				<div class='article-intro'>
					<p>[text intro]</p>
				</div>
				<h2>[subtitulo]</h2>
				<h2>[subtitulo]</h2>
				<p>{text parrafo}</p>
				<ul>(solo si es necesario usar lista)
					<li><strong>[caracteristica]</strong> [texto]</li>
					<li><strong>[caracteristica]</strong> [texto]</li>
					<li><strong>[caracteristica]</strong> [texto]</li>
				<ul>
				<h2>[subtitulo]</h2>
				<p>[texto parrafo]</p>
				<p>[texto parrafo]</p>
				<p>[texto parrafo]</p>
				<h2>URLs Sugeridas</h2>
				<div class='article-links'>
					<a href='[link]'>[link]</a>
					<a href='[link]'>[link]</a>
					<a href='[link]'>[link]</a>
				</a>
				<table></table>
				<div class='article-meta'>
					<p><strong>Meta Title: </strong>[meta title] | Hablando Claro</p>
					<p><strong>Meta Description: </strong>[meta description]</p>
				</div>"""
		print(f"====Urls sugeridas: {sugerencias_urls}")
		# --- 4. Construcción del Prompt (Optimizado) ---
		prompt = (
			f"Actúa como el **Redactor Senior de Productos y Ventas Digitales de Claro Perú**.\n"
			"Tu objetivo es crear un artículo útil, experto y altamente persuasivo que posicione a Claro como la mejor solución.\n\n"
			"Antes de redactar:\n"
			"- Analiza a profundidad la keyword y, si aplica, el tema específico.\n" # <--- Pequeño ajuste aquí
			"- Investiga temas, subtemas y preguntas relacionadas más buscadas en Google (autocomplete).\n"
			"- Integra esas búsquedas dentro del contenido de forma natural.\n"
			"- Usa redacción fluida, clara y natural, sin repetir innecesariamente la keyword.\n"
			"- Incluye sinónimos o términos relacionados estratégicamente.\n"
			"- Relaciona naturalmente conceptos con otros artículos o secciones del portal cuando sea posible.\n"
			"- Sugiere al lector, de manera útil y sutil, visitar páginas específicas del sitio web si quiere ampliar el tema.\n"
			"- Solo puedes redirigir a las URLs proporcionadas arriba, si el contenido lo permite de forma natural.\n"
			"- Cuando enlaces a otra URL, hazlo con frases como: 'como explicamos en nuestro artículo sobre...', 'puedes ver más en...', 'también te puede interesar...'.\n"
			"- Solo puedes mencionar la marca 'Claro' un máximo de dos veces en todo el artículo.\n\n"
			f"**DATOS DEL ARTÍCULO:**\n"
			f"- Keyword Principal: {keyword}\n"
			f"- Tema: {tema}\n"
			f"- Título Pre-aprobado: {title}\n"
			f"- URLs para sugerir al final: {sugerencias_urls}\n\n"

			f"{instruccion_producto}\n\n"

			f"**INFORMACIÓN DE BASE (Vertex AI):**\n{contexto_vertex}\n\n"

			"**REGLAS DE REDACCIÓN Y ESTILO (ESTRICTAS):**\n"
			"1. **CAPITALIZACIÓN (Sentence Case):** Está prohibido usar MAYÚSCULAS SOSTENIDAS en títulos<h1> o subtítulos<h2>. Usa siempre 'Tipo oración' (Ej: 'Internet en casa: Lo que necesitas saber'). Solo la primera letra va en mayúscula.\n"
			"2. **PUNTUACIÓN:** NUNCA termines un título o subtítulo con dos puntos (:), punto (.) o punto y coma (;).\n"
			"3. **MARCA:** Menciona 'Claro' de forma natural. Convierte características técnicas en beneficios emocionales.\n"
			"4. No capitalizes el texto en los titulo,subtitulos,h1 y h2 \n"
			"5. si vaz a poner '**' pon en su lugar <strong></strong> \n"
			"6. si vaz a generar una tabla usa <table><tr><td>\n"
			"7. **ENLACES:** Sugiere visitar otras secciones usando frases naturales como 'como te contamos en...', 'puedes ver más detalles en...'.\n\n"

			"**FORMATO HTML DE SALIDA (ESTRICTO):**\n"
			"**NUNCA DEVOLVER EL CÓDIGO HTML DENTRO DE BLOQUES MARKDOWN (NO USAR ```HTML).**\n" # <-- AÑADIR ESTA RESTRICCIÓN
			"**Devolver solo el mquetado html**\n"
			"Usa EXACTAMENTE estas etiquetas:\n\n"

			"<h1> <Escribe aquí el Título H1>\n"
			"Atractivo, CAPITALIZACIÓN (Sentence Case), sin dos puntos finales.</h1>\n\n"

			"<div class='article-intro'> <Párrafo introductorio si vas a poner ** reemplazalo por '<strong>'>\n"
			"Enganche (máx 4 líneas). La keyword principal ('" + keyword + "') debe estar en negrita.<div>\n\n"

			"<h2> <Primer Subtítulo H2>\n"
			"Sentence Case. Sin dos puntos.</h2>\n\n"

			"**Cuerpo del artículo:** Desarrolla el tema entre 600 y 1200 palabras. Todos los subtítulos dentro del cuerpo del artículo también deben estar en negrita (usando ****).\n"
			". Cada sección debe estar separada por saltos de línea (párrafos vacíos entre ellas).\n\n"
			"<p>\n"
			"en vez de '**' usa <h3>\n"
			"Desarrolla el contenido aquí. Usa párrafos cortos.\n"
			"- **NO uses viñetas (bullets).** Si enumeras características o listas (solo si es necesario sino el contenido que sea solo parrafo) usa <ul> <li>, en vez de '*' usa <ul> <li>, escribe: <strong>Característica:</strong> Descripción.\n"
			"- Si necesitas otro subtítulo, usa la etiqueta <h2> nuevamente, seguida de <p>.\n"
			"- Integra la marca Claro de forma natural.</p>\n\n"

			"Si hay URLs sugeridas muestra el title en <h2> URLs Sugeridas</h2>\n"
			"Si hay URLs sugeridas arriba, lístalas en <a> con su respectiva url:\n"


			"**coloca la seccion meta como meta_title y meta_description en <div class='article-meta'>\n"
			"Esta parte debe tener Máximo 65 caracteres debe terminar con '| Hablando Claro' y estar entre <p><strong><Escribe aquí el Meta title | Hablando Claro></strong></p>.\n"
			"La parte del meta description Máximo 160 caracteres. y entre <p><p>.\n"
			"la palabra Meta title y meta description no debe de estar soolo su contenido.\n"
			"Los h1,h2,h3 no debe contener strong.\n"
			"Una vez que generes el contenido, ordena el contenido si es necesario ubicalos en <ul> o <table>,\n"
			"si existe la palabra 'Meta title:' o 'Meta description:' removerlo.\n"
			"**si existe una palabra con '**' ponerlo entre <strong>\n"
			"antes de devolver,ordena el html,\n"
			f"Template base(no innecesariamente el la secuencia es asi):{template}\n"
		)

		# --- 5. Llamada a Gemini ---
		response = genai_client.models.generate_content(
			model="gemini-2.0-flash-001",
			contents=prompt
		)
		
		contenido_generado = response.text.strip()
		#Limpiar
		texto_limpio = contenido_generado.replace("```html\n", "", 1)
		texto_limpio = texto_limpio.rstrip().removesuffix("```")
		contenido_generado = texto_limpio.lstrip()

		# Validación simple
		if len(contenido_generado) < 50:
			 return "ERROR: Contenido insuficiente generado por la IA."

		return jsonify({
			"success": True,
			"texto": contenido_generado 
		})
	except Exception as e:
		error_msg = f"ERROR en generaContenidoIA: {e}"
		#print(error_msg)
		#return f"Error: {e}"
		return jsonify({
			"success": False,
			"texto": error_msg
		})
def get_json_tipo_contenido():
	base_dir = os.path.dirname(os.path.abspath(__file__))
	json_file_path = os.path.join(base_dir, '..', '..', '..', 'static', 'json', 'json_tipo-contenido.json')
	contenido_list = []
	try:
		with open(json_file_path, 'r', encoding='utf-8') as f:
			data = json.load(f) # Carga el JSON tal cual
		return data # Retorna el diccionario directamente
	except FileNotFoundError:
		print(f"Error: El archivo JSON no se encontró en la ruta: {json_file_path}")
		return []
	except json.JSONDecodeError:
		print(f"Error: No se pudo decodificar el archivo JSON '{json_file_path}'. Asegúrate de que sea un JSON válido.")
		return []
	except Exception as e:
		print(f"Ocurrió un error inesperado al cargar el JSON: {e}")
		return []
def get_urls_existentes():
	query = """
	SELECT * FROM prd-claro-mktg-data-storage.project_semantic_seo.hablando_claro_url
	"""
	return client.query(query).to_dataframe()
def elegir_urls_relevantes_con_gemini(keyword, urls_df):
	if urls_df.empty:
		return "No se encontraron URLs similares para sugerir."
	if 'page_url' not in urls_df.columns:
		return "Error interno: columna 'page_url' no encontrada."

	urls_texto = "\n".join([f"- {row['page_url']}" for _, row in urls_df.iterrows()])
	prompt = f"""
	Eres un experto en SEO. Tu tarea es identificar las 3 URLs más relevantes para una correcta implementación de enlaces internos.

	**Keyword Principal del Artículo:** '{keyword}'

	**Lista de URLs DISPONIBLES (SOLO puedes usar URLs de esta lista, NO inventes ni modifiques):**
	{urls_texto}

	**INSTRUCCIONES DE SALIDA ESTRICTAS:**
	Genera SOLO un listado numérico de 3 URLs. Cada URL debe estar en una nueva línea, precedida por su número y un punto, y sin ningún texto adicional o explicaciones.

	EJEMPLO DE FORMATO DE SALIDA (EXACTO):
	1. https://www.ejemplo.com/primera-url
	2. https://www.ejemplo.com/segunda-url
	3. https://www.ejemplo.com/tercera-url

	Asegúrate de que las URLs seleccionadas sean las más lógicamente relacionadas con la keyword principal.
	"""
	try:
		response = genai_client.models.generate_content(
			model="gemini-2.0-flash-001",
			contents=prompt
		)
		gemini_raw_response = response.text.strip()
		url_pattern = re.compile(r'^\d+\.\s*(https?://\S+)$', re.MULTILINE)
		suggested_urls = url_pattern.findall(gemini_raw_response)
		return suggested_urls if suggested_urls else "No se encontraron URLs relevantes para sugerir."
	except Exception as e:
		print(f"=====ERROR: elegir_urls_relevantes_con_gemini falló con una excepción: {e}")
		return "Error al generar sugerencias de URLs."
def get_urls_similares_por_keyword(keyword, urls_df):
	print(f"=====INFO: Buscando URLs similares para keyword: '{keyword}'")
	textos = urls_df["page_url"].tolist()
	# Asegúrate de que 'textos' no esté vacío antes de llamar a difflib.get_close_matches
	if not textos:
		print("ADVERTENCIA: No hay URLs en el DataFrame para buscar coincidencias.")
		return pd.DataFrame()
	matches = difflib.get_close_matches(keyword, textos, n=10, cutoff=0.2)
	print(f"INFO: Coincidencias de URLs encontradas: {matches}")
	return urls_df[urls_df["page_url"].isin(matches)].reset_index(drop=True)
def obtener_contexto_vertex(keyword):
	print(f"INFO: Obtener contexto de Vertex AI Search para keyword: '{keyword}'")
	project_id = "prd-claro-mktg-data-storage"
	location = "global"
	engine_id = "claro-peru_1748977843565"

	try:
		# Asegúrate de que todas las clases necesarias (ClientOptions, etc.) estén importadas en tu archivo.
		client_options = ClientOptions(api_endpoint=f"{location}-discoveryengine.googleapis.com") if location != "global" else None
		client = discoveryengine.ConversationalSearchServiceClient(client_options=client_options)

		serving_config = f"projects/{project_id}/locations/{location}/collections/default_collection/engines/{engine_id}/servingConfigs/default_serving_config"

		answer_generation_spec = discoveryengine.AnswerQueryRequest.AnswerGenerationSpec(
			ignore_adversarial_query=True,
			ignore_non_answer_seeking_query=True,
			ignore_low_relevant_content=False,
			model_spec=discoveryengine.AnswerQueryRequest.AnswerGenerationSpec.ModelSpec(model_version="stable"),
			prompt_spec=discoveryengine.AnswerQueryRequest.AnswerGenerationSpec.PromptSpec(
				preamble="Usa la información más relevante de tus fuentes indexadas, brinda el contexto necesario para usar info dentro de un articulo de Claro, comparte 3 urls relacionadas para que puedan buscar informacion mas detallada los usuarios."
			),
			include_citations=True,
			answer_language_code="es",
		)

		request = discoveryengine.AnswerQueryRequest(
			serving_config=serving_config,
			query=discoveryengine.Query(text=keyword),
			session=None,
			answer_generation_spec=answer_generation_spec,
		)

		print(f"INFO: Enviando solicitud a Vertex AI Search para keyword: '{keyword}'")
		response = client.answer_query(request)
		print(f"INFO: Respuesta de Vertex AI Search recibida. ¿Contiene respuesta?: {bool(response.answer)}")
		
		# Uso de acceso seguro (aunque ya lo tenías, lo mantenemos robusto)
		return response.answer.answer_text if response.answer and response.answer.answer_text else "No se pudo generar una respuesta de contexto relevante."
		
	except Exception as e:
		# 🛠️ CORRECCIÓN: Se elimina la línea 'st.error()' para evitar el NameError.
		# st.error(f"ERROR al obtener contexto de Vertex AI Search: {e}") 
		
		print(f"ERROR: obtener_contexto_vertex falló con: {e}")
		return f"Error al obtener contexto: {e}"
#Html Select json
@cltarticle_bp.route('select-tipo-contenido', methods=['GET', 'POST'])
def get_tipo_contenido(brand_id,proyecto_id):
	base_dir = os.path.dirname(os.path.abspath(__file__))
	#json_file_path = os.path.join(base_dir, '..', 'static', 'json', 'json_tipo-contenido.json')
	json_file_path = os.path.join(base_dir, '..', '..', '..', 'static', 'json', 'json_tipo-contenido.json')
	print(f"DEBUG: Intentando abrir el archivo en: {json_file_path}") 
	# ¡INICIALIZA la lista aquí!
	contenido_list = [] 
	try:
		with open(json_file_path, 'r', encoding='utf-8') as f:
			data = json.load(f)
		# Listar y construir la lista de diccionarios
		for tipo, detalles in data.items():
			contenido_item = {
				"nombre_tipo": tipo,
				"intencion": detalles.get("intencion", "N/A"),
				"descripcion": detalles.get("descripcion", "N/A"),
				"ejemplo": detalles.get("ejemplo", "N/A")
			}
			contenido_list.append(contenido_item)

		html_output = '<select id="cbo_tipo_articulo" name="cbo_tipo_articulo">\n'
		html_output += '\t<option value="">Selecciona una palabra</option>\n'
		for item in contenido_list:
			nombre = item["nombre_tipo"]
			descripcion = item["descripcion"]
			html_output += f'\t<option value="{nombre}" data_description="{descripcion}">{nombre}</option>\n'
		html_output += '</select>'
		return html_output
	except FileNotFoundError:
		print(f"Error: El archivo '{json_file_path}' no se encontró.")
		return "Error: Archivo de tipos de contenido no encontrado.", 500 
	except json.JSONDecodeError:
		print(f"Error: No se pudo decodificar el archivo JSON '{json_file_path}'. Asegúrate de que sea un JSON válido.")
		return "Error: Formato de archivo de tipos de contenido inválido.", 500
	except Exception as e:
		print(f"Ocurrió un error inesperado al cargar el JSON: {e}")
		return f"Error interno del servidor: {e}", 500

#Crear Documento HTML
def crear_html(texto_contenido, name_file):
	"""
	Crea un buffer BytesIO conteniendo el texto de entrada (texto_contenido)
	sin ninguna modificación, simulando la generación de un archivo .txt.
	"""
	# 1. Crear un buffer en memoria para el archivo.
	buffer = BytesIO()
	
	# 2. Escribir el texto de entrada (HTML/Texto) tal cual.
	# Es esencial codificar la cadena de texto (str) a bytes para escribir en BytesIO.
	buffer.write(texto_contenido.encode('utf-8'))
	
	# 3. Mover el puntero al inicio para que el contenido pueda ser leído.
	buffer.seek(0)
	
	return buffer

#Crear Documento Word
def crear_docx(texto_contenido, name_file):
	# -----------------------------------------------------
	# FASE 1: CONVERSIÓN de HTML a MARCADO INTERNO (CORREGIDA)
	# -----------------------------------------------------

	soup = BeautifulSoup(texto_contenido, 'lxml')
	custom_markup = []

	# Función auxiliar para convertir enlaces <a> a tu patrón [link=URL]TEXTO[/link]
	def process_links(element):
		html_segment = str(element)
		link_regex_html = re.compile(r'<a\s+href=["\'](.*?)["\']>(.*?)<\/a>', re.IGNORECASE)
		processed_text = link_regex_html.sub(r'[link=\1]\2[/link]', html_segment)
		
		clean_soup = BeautifulSoup(processed_text, 'lxml')
		return clean_soup.get_text()

	# Título Principal (H1)
	h1 = soup.find('h1')
	if h1:
		custom_markup.append(f"[title]{h1.get_text().strip()}")

	# Introducción (div.article-intro)
	intro_div = soup.find('div', class_='article-intro')
	if intro_div and intro_div.p:
		intro_text = process_links(intro_div.p)
		custom_markup.append(f"[intro]{intro_text.strip()}")
	
	# Variable de control
	urls_sugeridas_title_processed = False

	# Procesar el Cuerpo del Artículo
	body_elements = soup.find_all(['h2', 'h3', 'p', 'ul', 'ol', 'table', 'div'])
	
	for element in body_elements:
		tag_name = element.name
		
		# Saltar elementos ya procesados (H1 y INTRO)
		if element == h1 or (intro_div and element.find_parent('div', class_='article-intro') == intro_div):
			 continue

		# Títulos y Subtítulos (Manejo de H2/H3)
		elif tag_name in ['h2', 'h3']:
			title_text = element.get_text().strip()
			
			if "urls sugeridas" in title_text.lower():
				# 1. Si es "URLs Sugeridas", marcarlo como [title] para activar la bandera en FASE 2
				custom_markup.append(f"[title]{title_text}")
				urls_sugeridas_title_processed = True
			else:
				custom_markup.append(f"[subtitle]{title_text}")
		
		# Manejo de URLs Sugeridas (div.article-links)
		elif tag_name == 'div' and 'article-links' in element.get('class', []):
			if urls_sugeridas_title_processed:
				for a_tag in element.find_all('a'):
					url = a_tag.get('href')
					link_text = a_tag.get_text().strip()
					custom_markup.append(f"[item-link={url}]{link_text}")
			continue
			
		# Manejo de Metadatos (div.article-meta) - CORRECCIÓN DE DUPLICACIÓN
		elif tag_name == 'div' and 'article-meta' in element.get('class', []):
			custom_markup.append("[subtitle]Metadatos del Artículo")
			
			for p in element.find_all('p'):
				text_content = p.get_text().strip()
				
				# 1. Título Meta (Robusto con re.sub)
				if re.match(r'Meta Title:', text_content, re.IGNORECASE):
					# Usamos re.sub para limpiar el prefijo de forma robusta e ignorar mayúsculas/minúsculas
					cleaned_text = re.sub(r'Meta Title:\s*', '', text_content, flags=re.IGNORECASE).strip()
					custom_markup.append(f"[text]Título Meta: {cleaned_text}")
					
				# 2. Descripción Meta (Robusto con re.sub)
				elif re.match(r'Meta Description:', text_content, re.IGNORECASE):
					cleaned_text = re.sub(r'Meta Description:\s*', '', text_content, flags=re.IGNORECASE).strip()
					custom_markup.append(f"[text]Descripción Meta: {cleaned_text}")
					
				# 3. Si no es un meta tag conocido, se añade el texto original
				elif text_content: # Solo añadir si hay contenido y no es un meta tag conocido
					custom_markup.append(f"[text]{text_content}")
					
			continue

		# Párrafos (p)
		elif tag_name == 'p':
			text_content = process_links(element)
			if text_content.strip():
				custom_markup.append(f"[text]{text_content.strip()}")
			
		# Listas (UL y OL)
		elif tag_name in ['ul', 'ol']:
			for li in element.find_all('li', recursive=False):
				item_content = process_links(li)
				custom_markup.append(f"[text]* {item_content.strip()}")

		# Tablas (TABLE)
		elif tag_name == 'table':
			custom_markup.append("[subtitle]Contenido de Tabla")
			
			header_row = element.find('thead')
			if header_row:
				custom_markup.append("[text]ENCABEZADOS:")
				headers = [th.get_text().strip() for th in header_row.find_all('th')]
				custom_markup.append(f"[text]* {' | '.join(headers)}")
			
			body = element.find('tbody') or element
			for tr in body.find_all('tr'):
				cells = [td.get_text().strip() or th.get_text().strip() for td in tr.find_all(['td', 'th'])]
				if cells:
					custom_markup.append(f"[text]* {' | '.join(cells)}")

	# Obtener el contenido final en formato de marcado interno
	texto_contenido = '\n'.join(custom_markup)
	lines = texto_contenido.strip().split('\n')

	# -----------------------------------------------------
	# FASE 2: PROCESAMIENTO de MARCADO INTERNO (Sin cambios)
	# -----------------------------------------------------
	
	document = Document()
	inferred_keyword = name_file.replace(".docx", "").replace("_", " ").title()
	in_url_section = False
	
	for line in lines:
		line = line.strip()
		if not line:
			continue

		if line.startswith('[title]'):
			title_text = line.replace('[title]', '').strip()
			
			if "urls sugeridas" in title_text.lower():
				document.add_heading(title_text, level=3)
				in_url_section = True
			else:
				document.add_heading(title_text, level=1)
				in_url_section = False

		elif line.startswith('[intro]'):
			intro_text = line.replace('[intro]', '').strip()
			document.add_paragraph(intro_text, style='Intense Quote')

		elif line.startswith('[subtitle]'):
			subtitle_text = line.replace('[subtitle]', '').strip()
			document.add_heading(subtitle_text, level=3)
			in_url_section = False

		elif line.startswith('[text]'):
			text_content = line.replace('[text]', '').strip()
			
			# Detectar si es un ítem de lista (del procesamiento de UL/OL)
			is_list_item = text_content.startswith('* ')
			if is_list_item:
				text_content = text_content[2:].strip()
				style = 'List Bullet'
			else:
				style = None
			
			# Manejo de Enlaces (Link)
			link_regex = re.compile(r'\[link=(.*?)\](.*?)\[\/link\]')
			
			p = document.add_paragraph(style=style)
			last_idx = 0
			
			for match in link_regex.finditer(text_content):
				url = match.group(1)
				link_text = match.group(2)
				
				p.add_run(text_content[last_idx:match.start()])
				# Se llama a la función corregida
				add_hyperlink(p, url, link_text)
				last_idx = match.end()
			
			p.add_run(text_content[last_idx:])
			in_url_section = False

		elif line.startswith('[item-link='):
			if in_url_section:
				# Patrón: [item-link=URL_AQUI]TEXTO
				match = re.search(r'\[item-link=(.*?)\](.*)', line)
				if match:
					url = match.group(1)
					link_text = match.group(2).strip() or url
					
					p = document.add_paragraph(style='List Bullet')
					# Se llama a la función corregida
					add_hyperlink(p, url, link_text)
	
	# -----------------------------------------------------
	# FASE 3: Guardar y Devolver el Buffer
	# -----------------------------------------------------
	
	# 🌟 CORRECCIÓN DEL ERROR DE DATETIME
	# Cambiado de datetime.datetime.now() a datetime.now()
	fecha_doc = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
	
	document.add_page_break()
	document.add_heading("Metadatos de Generación", level=4)
	document.add_paragraph(f"Keyword: {inferred_keyword}")
	document.add_paragraph(f"Fecha de Generación: {fecha_doc}")

	buffer = BytesIO()
	document.save(buffer)
	buffer.seek(0)
	return buffer
	
def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
	"""
	Implementación alternativa más robusta de add_hyperlink.
	"""
	from docx.oxml.ns import qn
	from docx.oxml import OxmlElement
	from docx.opc.constants import RELATIONSHIP_TYPE as RT
	from docx.shared import RGBColor # Necesario para la corrección visual (asumiendo que está importado)
	
	# 1. Crear la relación externa (rId)
	rId = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)

	# 2. Crear el elemento w:hyperlink
	hyperlink = OxmlElement('w:hyperlink')
	hyperlink.set(qn('r:id'), rId)
	
	# 3. Crear el elemento w:r (Run) dentro del hipervínculo
	new_run = OxmlElement('w:r')
	
	# 4. APLICACIÓN DE ESTILO CRÍTICA (w:rPr para subrayado y color)
	# Se debe crear y configurar el XML de las propiedades del Run
	rPr = OxmlElement('w:rPr')
	
	# a. Subrayado
	if underline:
		u = OxmlElement('w:u')
		u.set(qn('w:val'), 'single')
		rPr.append(u)

	# b. Color (Para que se vea azul como un enlace)
	c = OxmlElement('w:color')
	c.set(qn('w:val'), color)
	rPr.append(c)

	new_run.append(rPr) # Añadir las propiedades al Run
	
	# 5. Crear w:t (Text) dentro del Run
	text_element = OxmlElement('w:t')
	text_element.text = text
	new_run.append(text_element)
	
	# 6. Agregar el Run al Hyperlink
	hyperlink.append(new_run)
	
	# 7. Agregar el Hyperlink al Párrafo
	paragraph._p.append(hyperlink)

	# ELIMINAR EL CÓDIGO FALLIDO DE APLICACIÓN DE FORMATO
	# Ya no es necesario el bloque original de aplicación de estilo.
	return new_run # Retornar el elemento XML creado.

@cltarticle_bp.route('download-generated-article', methods=['POST'])
def download_generated_article(brand_id, proyecto_id):
	# Recuperar el contenido guardado en la sesión
	texto_plain_for_docx = request.values.get('content')
	keyword_for_docx = request.values.get('keyword')
	data_titulo = request.values.get('titulo')
	if not texto_plain_for_docx or not keyword_for_docx:
		# Si no hay contenido en la sesión, redirigir o mostrar un error
		return "No hay artículo para descargar. Por favor, genera un artículo primero.", 400
	name_file = f"{keyword_for_docx}_articulo.docx"
	# Llamar a crear_docx con los datos recuperados de la sesión
	docx_buffer = crear_docx(texto_plain_for_docx, name_file)
	# Retornar el archivo como una respuesta de Flask
	return send_file(
		docx_buffer,
		mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
		as_attachment=True,
		download_name=name_file
	)

@cltarticle_bp.route('download-html', methods=['POST'])
def download_generated_html(brand_id, proyecto_id):
	# Recuperar el contenido guardado en la sesión
	texto_plain_for_docx = request.values.get('content')
	keyword_for_docx = request.values.get('keyword')
	data_titulo = request.values.get('titulo')
	if not texto_plain_for_docx or not keyword_for_docx:
		# Si no hay contenido en la sesión, redirigir o mostrar un error
		return "No hay artículo para descargar. Por favor, genera un artículo primero.", 400
	name_file = f"{keyword_for_docx}_articulo.html"
	# Llamar a crear_docx con los datos recuperados de la sesión
	docx_buffer = crear_html(texto_plain_for_docx, name_file)
	# Retornar el archivo como una respuesta de Flask
	return send_file(
		docx_buffer,
		mimetype="text/plain", 
		as_attachment=True,
		download_name=name_file
	)



@cltarticle_bp.route('generated-image', methods=["GET", "POST"])
def generateImageIA(brand_id,proyecto_id):
	data_prompt = request.args.get('promt', '')
	data_bytes, mime_type = generar_imagen_imagen4(data_prompt)

	if data_bytes is None:
		return jsonify({"error": "Fallo al generar la imagen.", "details": "El modelo no devolvió datos."}), 500

	try:
		# 3. Guardar en Google Cloud Storage
		folder_path = "blog_cover" # Un identificador único para el nombre de archivo
		
		# 'name_fileext' es el nombre del archivo dentro del bucket (ej: 20251106_gamefest_banner_abc123.png)
		name_fileext = save_image_to_gcs(
			data=data_bytes, 
			mime_type=mime_type, 
			folder_path=folder_path
		)

		# 4. Generar URL Firmada
		# La función generar_url_firmada requiere el nombre del archivo (blob_name)
		url_firmada = generar_url_firmada(
			bucket_name=bucket_name,
			blob_name=name_fileext # Pasamos solo el nombre del archivo
		)
		# 4. Save
		imageia_id = generarCodigo("IA");
		imgia_name = name_fileext
		imgia_promt = data_prompt
		imgia_type = "blog"
		imgia_fecha = fechaActual()
		proyecto_id = request.values.get('txt_pid', '')

		new_data = Imagesia(
			imgia_id = imageia_id,
			imgia_name = imgia_name,
			imgia_promt = imgia_promt,
			imgia_type = imgia_type,
			imgia_fecha = imgia_fecha,
			imgia_estado = 1,
			proyecto_id = proyecto_id
		)
		db.session.add(new_data)
		db.session.commit()
		# 5. Devolver la URL
		return jsonify({
			"success": True,
			"message": "Imagen generada y guardada en GCS.",
			"imagen_temporal": url_firmada,
			"nombre_archivo": name_fileext
		})

	except Exception as e:
		# Esto captura errores de GCS (conexión, permisos, etc.)
		return jsonify({"error": f"Error al guardar o firmar la URL: {str(e)}"}), 500

@cltarticle_bp.route('generated-promt', methods=["GET"])
def generatePromtImageIA(brand_id, proyecto_id):
	# 1. RECUPERAR PARÁMETROS DE LA SOLICITUD GET
	# 'request' debe estar importado y disponible aquí.
	tipo_articulo = request.args.get('tipo', '')
	tipo_description = request.args.get('description', '') # Variable no usada en el prompt final, pero mantenida.
	tema_articulo = request.args.get('tema', '')
	keyword_articulo = request.args.get('keyword', '')
	titulo_articulo = request.args.get('titulo', '') 
	motivo_articulo = request.args.get('motivo', '')
	
	# Parámetros Fijos o Base
	render_calidad = 'Render 8K'
	
	# 2. Lógica para Estilo y Tono Emocional (DEBE IR ANTES DE CONSTRUIR EL PROMPT)
	estilo_visual = "clean, modern, premium style" 
	tono_emocional = "Calma y Seguridad" # Tono por defecto
	
	if tipo_articulo.lower() == "educativo":
		estilo_visual = "minimalist, informative, friendly visual style"
		tono_emocional = "Confianza, Claridad y Aprendizaje"
	elif tipo_articulo.lower() == "guía":
		estilo_visual = "detailed, structured, high-quality rendering style"
		tono_emocional = "Organización, Eficiencia y Logro"

	# 3. PREPARAR CLÁUSULAS (Asegurando que la definición precede al uso)

	# A1. Concepto principal del tema
	concepto_principal = f"Scene Concept: {tema_articulo}, emphasizing the seamless use of the {keyword_articulo} service."
	
	# A2. Elementos visuales y restricción suave
	elementos_visuales = f"Visual Elements: Subtly incorporate discrete visual elements related to the keyword '{keyword_articulo}'"
	
	# A3. Lógica correcta para Motivo y Enfoque (DEBE RESOLVERSE CON IF/ELSE)
	if motivo_articulo:
		motivo_enfoque_contenido = motivo_articulo
	else:
		motivo_enfoque_contenido = "Familia tradicional (padres hombre y mujer) en casa usando dispositivos conectados."
		
	motivo_enfoque = f"Image theme and focus: the image is generated for: {motivo_enfoque_contenido}"


	# A4. Lógica de Fondo (Background/Setting) - Anti-Borde Blanco
	background_setting_base = f"If the keyword '{keyword_articulo}' refers to a specific place, set it as the background using a modern and premium filter. Otherwise, use a modern and warm home living room setting."
	background_setting = f"Background/Setting: {background_setting_base}, seamless background, image must occupy the entire viewport."
	
	# B. DETALLES DE COMPOSICIÓN Y TÉCNICOS
	# El tono emocional ya se ha resuelto en el punto 2.
	adicionales = [
		f"Full-frame render, Wide-Screen aspect ratio, Cinematic", 
		f"Artistic Style: realistic",
		f"Lighting: Studio light soft",
		f"Camera Angle: Normal (Eye-level)",
		f"Emotional Tone: {tono_emocional}", # Se usa la variable resuelta
		f"Composition Rule: Rule of Thirds",
		f"Depth of Field: 50 (Bokeh effect)."
	]
	
	clausulas_adicionales_str = ", ".join(adicionales)
	clausulas_adicionales_final = f"Composition and Technical Details: {clausulas_adicionales_str}"

	# C. INSTRUCCIONES CRÍTICAS REFORZADAS
	# Se condensa en una sola línea para evitar problemas de formato.
	critical_instructions = (
		"ULTRA CRITICAL: ABSOLUTELY NO TEXT, NO LOGOS, NO WATERMARKS, NO UI CHROME, NO DEBUG ARTIFACTS, NO BOUNDING BOXES, "
		"NO black bars cinema, NO ANNOTATIONS, NO GUIDES, AND NO WIFI/SIGNAL ICONS ARE PERMITTED ANYWHERE."
		"NO extra limbs, NO mutated hands, NO deformed fingers, NO extra arms. PURE IMAGE OUTPUT ONLY."
	)
	
	# 4. ENSAMBLAJE FINAL DEL PROMPT (Usando f-strings de forma limpia)
	mi_prompt = (
			f"Generame y arregla my promt que es el siguiente:\n"
			f"Generate an {render_calidad}, full-frame, cinematic image for telco.\n"
			f"Visual style: {estilo_visual}, realistic. ubica los elementos de forma coherente.\n"
			f"{concepto_principal}\n"
			f"{elementos_visuales}\n"
			f"{motivo_enfoque}\n" # Se inserta la variable ya resuelta
			f"{background_setting}\n"
			f"{clausulas_adicionales_final}\n"
			f"{critical_instructions}"
	)
	
	# Se eliminan los saltos de línea y espacios excesivos
	mi_prompt = mi_prompt.strip()
	final_prompt = generarPrompIA(mi_prompt)
	# --- FIN DE LA LÓGICA DE CONSTRUCCIÓN DEL PROMPT ---
	
	# 5. DEVOLVER EL PROMPT EN LA RESPUESTA JSON
	return jsonify({
		"success": True,
		"promt": final_prompt
	})


def generar_imagen_imagen4(prompt: str):
	try:
		# Aquí debes asegurarte de que la clave API esté disponible
		# Nota: Es altamente inseguro dejar la clave API dura en el código. Úsala solo para desarrollo/prueba.
		clave_secreta_gemini = os.environ.get("GEMINI_API_KEY", "AIzaSyBjLIPcIzCY6zuWJdiikno0sWHpilbwLw4")
		if not clave_secreta_gemini:
			print("Error: Clave GEMINI_API_KEY no configurada.", file=sys.stderr)
			return None, None
			
		client = genai.Client(api_key=clave_secreta_gemini)
		
		response = client.models.generate_images(
			model='imagen-4.0-generate-001',
			prompt=prompt,
			config=types.GenerateImagesConfig(
				number_of_images=1,
				#  CAMBIO 1: Usamos el parámetro aspect_ratio en la configuración
				aspect_ratio='16:9',
			)
		)

		if response.generated_images:
			img_bytes = response.generated_images[0].image.image_bytes
			mime_type = 'image/png'
			
			#  CAMBIO 2: Devuelve los bytes y el tipo MIME (para GCS)
			return img_bytes, mime_type
		else:
			print("No se pudo generar la imagen. La respuesta no contiene imágenes generadas.", file=sys.stderr)
			return None, None

	except Exception as e:
		print(f"Ocurrió un error en la generación: {e}", file=sys.stderr)
		return None, None

# --- FUNCIONES AUXILIARES (Sin Cambios Relevantes) ---
def save_image_to_gcs(data, mime_type, folder_path) -> str:
	if isinstance(data, str):
		data_bytes = base64.b64decode(data)
	else:
		data_bytes = data
	ext = mimetypes.guess_extension(mime_type)
	ext = ext.lstrip('.') if ext else "png" 
	fecha_utc = datetime.datetime.utcnow().strftime('%Y%m%d')
	name_archivo = f"{fecha_utc}_{folder_path}_{uuid.uuid4().hex}"
	name_fileext = f"{name_archivo}.{ext}"
	nombre_archivo = f"generate/blog/{name_fileext}" 
	bucket = storage_client.bucket(bucket_name)
	blob = bucket.blob(nombre_archivo)
	blob.upload_from_string(data_bytes, content_type=mime_type)
	return name_fileext


#------------------------
#-------Funciones
#------------------------

def generar_url_firmada(bucket_name, blob_name):
	bucket = storage_client.bucket(bucket_name)
	blob_path = f"generate/blog/{blob_name}"
	blob = bucket.blob(blob_path)
	url = blob.generate_signed_url(
		version="v4",
		expiration=timedelta(minutes=15),
		method="GET",
	)
	#url = blob.generate_signed_url(expiration=3600)
	return url
def listarImages(data_id):
	try:
		query = Imagesia.query.filter_by(imgia_type='blog',imgia_estado=1,proyecto_id=data_id)
		lista = query.all()
		return lista
	except Exception as e:
		return jsonify({"error": str(e)}), 500

def generarPrompIA(prompt):
	"""
	Envía un prompt a la IA para su mejora y extrae el resultado más limpio.
	"""
	try:
		# Usando la línea original, asumiendo que 'genai_client' está definido
		# NOTA: genai_client debe ser una variable global o pasada como argumento si no está
		# definida en este scope. Asumo que está definida globalmente por el import de config.
		response = genai_client.models.generate_content(
			model="gemini-2.0-flash-001",
			contents=prompt
		)
		contenido_generado = response.text
		
		# Se elimina la verificación de longitud de < 5
		return contenido_generado
		
	except Exception as e:
		error_msg = f"ERROR: Fallo crítico en la función generarPrompIA: {e}"
		# Se usa logging.error, asumiendo que logging está importado
		logging.error(error_msg)  
		
		# Retornamos el error como el prompt
		return f"Error de ejecución al mejorar el prompt: {e}"